#!/usr/bin/env python3
"""
Generate a .docx release notes document from a JSON data file.

Usage:
    python3 generate_docx.py <input.json> [output.docx]

The JSON file should have this structure:
{
    "version": "0.54.0",
    "title": "Focus & Flow",
    "sections": [
        {
            "heading": "Features",
            "items": [
                {
                    "id": "JS-292",
                    "text": "Note text. Can contain **bold** markers."
                }
            ]
        },
        {
            "heading": "Quality of Life Improvements",
            "items": [...]
        },
        {
            "heading": "Bug Fixes",
            "items": [...]
        }
    ]
}

Each item's "id" becomes a hyperlink to https://linear.app/anytype/issue/<id>.
Items appear in the order provided (should match Linear ordering).
Text supports **bold** markers which become bold runs in the docx.
"""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


LINEAR_BASE = "https://linear.app/anytype/issue"


def add_hyperlink(paragraph, url, text, font_size=None, bold=False):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "4A90D9")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    if font_size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(font_size * 2))
        rPr.append(sz)

    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)

    run.append(rPr)
    run.text = text
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    return hyperlink


def add_rich_text(paragraph, text, base_size=None):
    """
    Add text to a paragraph, handling **bold** markers.
    Returns nothing; modifies paragraph in place.
    """
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if base_size:
                run.font.size = Pt(base_size)
        else:
            if part:
                run = paragraph.add_run(part)
                if base_size:
                    run.font.size = Pt(base_size)


def generate_docx(data, output_path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # Title
    version = data["version"]
    title_text = data.get("title", "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Release {version}")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    if title_text:
        run = p.add_run(f"  {title_text}")
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    for section in data["sections"]:
        heading = section["heading"]
        items = section.get("items", [])

        if not items:
            continue

        # Section heading
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(heading)
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        for item in items:
            issue_id = item.get("id", "")
            text = item.get("text", "")

            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(0)

            if issue_id:
                issue_url = f"{LINEAR_BASE}/{issue_id}"
                add_hyperlink(p, issue_url, issue_id, font_size=10, bold=True)
                run = p.add_run("  ")
                run.font.size = Pt(11)

            # Note text with bold support
            if text:
                add_rich_text(p, text, base_size=11)

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 2:
        print(f"Usage: {sys.argv[0]} <input.json> [output.docx]", file=sys.stderr)
        sys.exit(1)

    input_path = Path(sys.argv[1])
    if not input_path.exists():
        print(f"Error: {input_path} not found", file=sys.stderr)
        sys.exit(1)

    with open(input_path) as f:
        data = json.load(f)

    version = data.get("version", "unknown")
    default_output = input_path.parent / f"release-{version}.docx"
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else default_output

    result = generate_docx(data, str(output_path))
    print(f"Generated: {result}")


if __name__ == "__main__":
    main()
